"""Forensic reporting (Phase 7, FR-16/17).

Generates an investigation-ready report (PDF via reportlab, Word via python-docx) with:
executive summary, top risk entities + reasons, correlation evidence (with provenance),
money-flow findings, and an optional Suspicious Transaction Report (STR) section.

Every finding carries provenance / references so the output is evidentiary (NFR-7).
"""

from __future__ import annotations

import dataclasses
import os
from datetime import datetime, timedelta, timezone

from ..detection import service as detection

IST = timezone(timedelta(hours=5, minutes=30))


def _stamp() -> str:
    # Wall-clock is fine here (report generation, not part of the deterministic pipeline).
    return datetime.now(IST).strftime("%Y%m%d_%H%M%S")


def _top_entities(data, n=15):
    """Highest risk first, on the same ranking key the API serves — see `detection.risk_rank`.

    The report and the screen must not disagree about who is worst; a second copy of the
    ordering rule here is how that starts.
    """
    return sorted(data["risk"].values(), key=detection.risk_rank)[:n]


def _top_transfers(data, n=15):
    return sorted(data["transfers"], key=lambda t: -(t.get("amount") or 0))[:n]


def _build_charts(data, out_dir) -> list[tuple[str, str]]:
    """E2: render charts (risk bars, activity timeline) as PNGs for the report."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    charts = []
    stamp = _stamp()

    # 1. Top risk entities bar
    top = [r for r in _top_entities(data, 10) if r["risk_score"] > 0]
    if top:
        fig, ax = plt.subplots(figsize=(7, 3.2))
        labels = [str(r["label"])[:18] for r in top][::-1]
        vals = [r["risk_score"] for r in top][::-1]
        ax.barh(labels, vals, color="#c0392b")
        ax.set_xlabel("Risk score")
        ax.set_title("Top risk entities")
        fig.tight_layout()
        p = os.path.join(out_dir, f"chart_risk_{stamp}.png")
        fig.savefig(p, dpi=120)
        plt.close(fig)
        charts.append(("Top risk entities", p))

    # 2. Activity-over-time timeline (events per day by type)
    from collections import defaultdict
    by_day = defaultdict(lambda: defaultdict(int))
    for e in data.get("events", []):
        ts = e.get("timestamp_start")
        if ts:
            by_day[str(ts)[:10]][e["event_type"]] += 1
    if by_day:
        days = sorted(by_day)[:120]
        fig, ax = plt.subplots(figsize=(7, 3.0))
        for et, color in (("TRANSACTION", "#d62728"), ("CALL", "#1f77b4"),
                          ("IP_SESSION", "#2ca02c")):
            ax.plot(days, [by_day[d].get(et, 0) for d in days], label=et, color=color)
        ax.set_title("Activity over time")
        ax.legend(fontsize=7)
        ax.set_xticks(days[:: max(1, len(days) // 8)])
        ax.tick_params(axis="x", labelrotation=45, labelsize=6)
        fig.tight_layout()
        p = os.path.join(out_dir, f"chart_timeline_{stamp}.png")
        fig.savefig(p, dpi=120)
        plt.close(fig)
        charts.append(("Activity over time", p))

    return charts


def payload_from_investigation(inv, dataset: str, window_minutes: int) -> dict:
    """Build the report payload from an Investigation.

    `generate` needs `summary` and `window`, neither of which is a field on the
    Investigation dataclass — so every caller had to know to add them, and the
    failure mode was a bare `KeyError: 'window'` from deep inside PDF assembly.
    Keeping that assembly here means one place knows the contract.
    """
    data = {f.name: getattr(inv, f.name) for f in dataclasses.fields(inv)}
    data["summary"] = inv.summary()
    data["window"] = window_minutes
    data["dataset"] = dataset
    # The report prints os.path.basename(input_dir) as the case name; for an uploaded
    # or renamed dataset that is the more meaningful label.
    data.setdefault("input_dir", dataset)
    return data


def generate(data: dict, out_dir: str, fmt: str = "pdf") -> str:
    os.makedirs(out_dir, exist_ok=True)
    if fmt == "docx":
        return _generate_docx(data, out_dir)
    return _generate_pdf(data, out_dir)


# --------------------------------------------------------------------------- PDF
def _generate_pdf(data, out_dir) -> str:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    path = os.path.join(out_dir, f"forensic_report_{_stamp()}.pdf")
    doc = SimpleDocTemplate(path, pagesize=A4, topMargin=15 * mm, bottomMargin=15 * mm)
    styles = getSampleStyleSheet()
    S = data["summary"]
    story = []

    story.append(Paragraph("ERakshak — Forensic Investigation Report", styles["Title"]))
    story.append(Paragraph("AI-Powered Financial &amp; Telecom Dataset Analyzer "
                           "(Bank · CDR · IPDR Fusion)", styles["Normal"]))
    story.append(Paragraph(f"Generated: {datetime.now(IST).strftime('%Y-%m-%d %H:%M IST')} · "
                           f"Dataset: {os.path.basename(data['input_dir'])} · "
                           f"Window W = {data['window']} min", styles["Normal"]))
    story.append(Spacer(1, 10))

    story.append(Paragraph("1. Executive Summary", styles["Heading2"]))
    summ = [["Files", S["files"]], ["Events", S["events"]],
            ["Transactions / Calls / IP sessions",
             f"{S['transactions']} / {S['calls']} / {S['ip_sessions']}"],
            ["Resolved entities", S["entities"]],
            ["Correlation coincidences", S["correlation_hits"]],
            ["High-risk entities", S["high_risk_entities"]],
            ["Rejected rows", S["rejected_rows"]]]
    story.append(_pdf_table(summ, Table, TableStyle, colors, header=False))
    story.append(Spacer(1, 10))

    # E2: charts
    for title, img in _build_charts(data, out_dir):
        story.append(Paragraph(title, styles["Heading3"]))
        story.append(Image(img, width=170 * mm, height=70 * mm))
        story.append(Spacer(1, 8))

    story.append(Paragraph("2. Top Risk Entities", styles["Heading2"]))
    rows = [["Entity", "Risk", "Band", "Reasons"]]
    for r in _top_entities(data):
        rows.append([str(r["label"])[:28], r["risk_score"], r["band"],
                     ", ".join(sorted({f["rule"] for f in r["rule_flags"]}))[:60]])
    story.append(_pdf_table(rows, Table, TableStyle, colors))
    story.append(Spacer(1, 10))

    story.append(Paragraph("3. Cross-Dataset Correlation Evidence", styles["Heading2"]))
    if not data["correlation_hits"]:
        story.append(Paragraph("No call+IP+transfer coincidences at the current window.",
                               styles["Normal"]))
    for h in sorted(data["correlation_hits"],
                    key=lambda x: -(x["transaction"].get("amount") or 0))[:12]:
        story.append(Paragraph(f"• <b>{h['entity_label']}</b>: {h['explanation']}",
                               styles["Normal"]))
        prov = h["transaction"].get("provenance", {})
        story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;<i>Evidence: ref {h['transaction'].get('ref_no')} · "
                               f"source {prov.get('source_file')} row {prov.get('row')}</i>",
                               styles["Normal"]))
    story.append(Spacer(1, 10))

    story.append(Paragraph("4. Money-Flow Findings (largest transfers)", styles["Heading2"]))
    rows = [["From", "To", "Amount", "Ref"]]
    for t in _top_transfers(data):
        rows.append([_label(data, t["from_entity"])[:22], _label(data, t["to_entity"])[:22],
                     t.get("amount"), t.get("ref")])
    story.append(_pdf_table(rows, Table, TableStyle, colors))
    story.append(Spacer(1, 10))

    story.append(Paragraph("5. Suspicious Transaction Report (STR) — Draft", styles["Heading2"]))
    for line in _str_lines(data):
        story.append(Paragraph(line, styles["Normal"]))

    rows = _eligibility_rows(data)
    if rows:
        story.append(Paragraph("6. Detection Audit — which rules could run", styles["Heading2"]))
        story.append(Paragraph(
            "A rule that found nothing and a rule that could not run are different findings. "
            "`eligible` is how many candidates the rule had in this case; `fired` is how many "
            "it flagged. Zero eligible is a statement about the evidence, not a miss.",
            styles["Normal"]))
        story.append(Spacer(1, 4))
        # Fixed widths so the Note column has somewhere to wrap into; without them reportlab
        # sizes to the longest unbroken string and pushes the last column off the page.
        story.append(_pdf_table(rows, Table, TableStyle, colors, wrap_col=4,
                                col_widths=[104, 44, 50, 34, 218]))

    story.append(Spacer(1, 12))
    story.append(Paragraph("<i>Generated by ERakshak. Findings are investigative leads derived "
                           "from the supplied data and require analyst verification.</i>",
                           styles["Italic"]))
    doc.build(story)
    return path


def _pdf_table(rows, Table, TableStyle, colors, header=True, wrap_col=None, col_widths=None):
    """Render rows as a PDF table.

    `wrap_col` flows that column's text as a Paragraph instead of a single unbreakable string.
    A plain reportlab cell does not wrap, so long prose runs off the page edge — which is why
    the eligibility note used to be cut at 96 characters. Truncating the sentence that explains
    why a rule could not run defeats the point of printing it, and silently dropping text is
    the one thing rule 2 forbids; wrapping keeps all of it on the page.
    """
    body = [[str(c) for c in r] for r in rows]
    if wrap_col is not None:
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import Paragraph as _P
        cell = ParagraphStyle("cell", fontName="Helvetica", fontSize=8, leading=9.5)
        for r in body[1:] if header else body:
            if r[wrap_col]:
                r[wrap_col] = _P(r[wrap_col], cell)
    t = Table(body, repeatRows=1 if header else 0, colWidths=col_widths)
    style = [("FONTSIZE", (0, 0), (-1, -1), 8),
             ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
             ("VALIGN", (0, 0), (-1, -1), "TOP")]
    if header:
        style += [("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
                  ("TEXTCOLOR", (0, 0), (-1, 0), colors.white)]
    t.setStyle(TableStyle(style))
    return t


# -------------------------------------------------------------------------- DOCX
def _generate_docx(data, out_dir) -> str:
    from docx import Document

    path = os.path.join(out_dir, f"forensic_report_{_stamp()}.docx")
    doc = Document()
    S = data["summary"]
    doc.add_heading("ERakshak — Forensic Investigation Report", 0)
    doc.add_paragraph("AI-Powered Financial & Telecom Dataset Analyzer (Bank · CDR · IPDR Fusion)")
    doc.add_paragraph(f"Generated: {datetime.now(IST).strftime('%Y-%m-%d %H:%M IST')} · "
                      f"Dataset: {os.path.basename(data['input_dir'])} · Window W = {data['window']} min")

    doc.add_heading("1. Executive Summary", 1)
    for k, v in [("Files", S["files"]), ("Events", S["events"]),
                 ("Transactions/Calls/IP", f"{S['transactions']}/{S['calls']}/{S['ip_sessions']}"),
                 ("Entities", S["entities"]), ("Correlation hits", S["correlation_hits"]),
                 ("High-risk entities", S["high_risk_entities"]), ("Rejected rows", S["rejected_rows"])]:
        doc.add_paragraph(f"{k}: {v}", style="List Bullet")

    # E2: charts
    from docx.shared import Inches
    for title, img in _build_charts(data, out_dir):
        doc.add_heading(title, 3)
        doc.add_picture(img, width=Inches(6.2))

    doc.add_heading("2. Top Risk Entities", 1)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(["Entity", "Risk", "Band", "Reasons"]):
        tbl.rows[0].cells[i].text = h
    for r in _top_entities(data):
        c = tbl.add_row().cells
        c[0].text = str(r["label"])
        c[1].text = str(r["risk_score"])
        c[2].text = r["band"]
        c[3].text = ", ".join(sorted({f["rule"] for f in r["rule_flags"]}))

    doc.add_heading("3. Cross-Dataset Correlation Evidence", 1)
    if not data["correlation_hits"]:
        doc.add_paragraph("No call+IP+transfer coincidences at the current window.")
    for h in sorted(data["correlation_hits"],
                    key=lambda x: -(x["transaction"].get("amount") or 0))[:12]:
        doc.add_paragraph(f"{h['entity_label']}: {h['explanation']}", style="List Bullet")

    doc.add_heading("4. Money-Flow Findings", 1)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(["From", "To", "Amount", "Ref"]):
        tbl.rows[0].cells[i].text = h
    for t in _top_transfers(data):
        c = tbl.add_row().cells
        c[0].text = _label(data, t["from_entity"])
        c[1].text = _label(data, t["to_entity"])
        c[2].text = str(t.get("amount"))
        c[3].text = str(t.get("ref"))

    doc.add_heading("5. Suspicious Transaction Report (STR) — Draft", 1)
    for line in _str_lines(data):
        doc.add_paragraph(line, style="List Bullet")

    rows = _eligibility_rows(data)
    if rows:
        doc.add_heading("6. Detection Audit — which rules could run", 1)
        doc.add_paragraph(
            "A rule that found nothing and a rule that could not run are different findings. "
            "Zero eligible is a statement about the evidence, not a missed detection.")
        t = doc.add_table(rows=1, cols=len(rows[0]))
        t.style = "Light Grid Accent 1"
        for i, head in enumerate(rows[0]):
            t.rows[0].cells[i].text = str(head)
        for r in rows[1:]:
            cells = t.add_row().cells
            for i, v in enumerate(r):
                cells[i].text = str(v)

    doc.add_paragraph("Generated by ERakshak. Findings require analyst verification.")
    doc.save(path)
    return path


def _eligibility_rows(data) -> list[list]:
    """Per-rule audit table for the report, or [] when the pipeline did not produce one.

    Included because `0 high-risk entities` is the headline an investigator acts on, and
    without this table there is no way to tell a clean case from an inert detector.
    """
    rows = data.get("rule_eligibility") or []
    if not rows:
        return []
    out: list[list] = [["Rule", "Enabled", "Eligible", "Fired", "Note"]]
    for r in rows:
        eligible = r.get("eligible")
        out.append([
            r.get("rule"),
            "yes" if r.get("enabled") else "no",
            "n/a" if eligible is None else f"{eligible:,}",
            r.get("fired", 0),
            # In full. This was cut at 96 characters, which decapitated every inert note —
            # "no entity is seen both receiving and sending, so forwarding cannot be observed
            # — a one-hop view of the money trail" is 113 — leaving the reader the premise
            # and not the conclusion. The PDF renderer wraps this column instead.
            r.get("note") or "",
        ])
    return out


def _label(data, eid):
    return str((data["entities"].get(eid, {}) or {}).get("label", eid))


#: Grounds and basis clauses printed per entry before summarising the rest. Both were
#: silently truncated at 5 and 3, so an entity with more typologies lost the extras with no
#: indication — the one thing rule 2 forbids, in the document most likely to be relied on.
_STR_MAX_GROUNDS = 5
_STR_MAX_BASIS = 3


def _str_subject(data, row) -> str:
    """Identify the subject by typed identifier, not just a label.

    `label` on real data is a bare account number or phone, so "Suspected subject:
    50100369668648" leaves the reader guessing what kind of identifier that is. An STR has
    to name the account it concerns.
    """
    eid = row.get("entity_id")
    ent = (data.get("entities") or {}).get(eid) or {}
    ids = ent.get("identifiers") or set()
    typed = sorted({f"{t}:{v}" for (t, v) in ids
                    if t in ("ACCOUNT_NO", "PHONE", "UPI_ID", "IMEI", "IMSI")})
    shown = ", ".join(typed[:4])
    if len(typed) > 4:
        shown += f" (+{len(typed) - 4} more)"
    label = row.get("label") or eid or "unidentified"
    return f"{label} [{shown}]" if shown else str(label)


def _str_lines(data) -> list[str]:
    """FIU-IND-style STR narrative lines, one suspected subject per entry.

    This is a **draft for analyst review**, and the wording says so rather than asserting a
    filing decision. The recommended action is graded by band: a high-band entity is put
    forward for filing, a medium-band one for review. Recommending a freeze on an
    automatically-scored medium entity — the previous behaviour, which fired down to a score
    of 48.6 — is an action against a real account holder that a rules engine should not be
    proposing on its own.
    """
    lines = []
    n = 0
    for r in _top_entities(data, 10):
        if r["band"] not in ("high", "medium") or not r["rule_flags"]:
            continue
        n += 1
        flags = r["rule_flags"]
        grounds = "; ".join(f["rule"].replace("_", " ") for f in flags[:_STR_MAX_GROUNDS])
        if len(flags) > _STR_MAX_GROUNDS:
            grounds += f" (+{len(flags) - _STR_MAX_GROUNDS} further typolog" \
                       f"{'y' if len(flags) - _STR_MAX_GROUNDS == 1 else 'ies'})"
        basis = "; ".join(f["detail"] for f in flags[:_STR_MAX_BASIS])
        if len(flags) > _STR_MAX_BASIS:
            basis += f" (+{len(flags) - _STR_MAX_BASIS} further ground" \
                     f"{'' if len(flags) - _STR_MAX_BASIS == 1 else 's'} on file)"

        feats = r.get("features") or {}
        txns = int(feats.get("txn_count") or 0)
        particulars = (
            f"{txns} transaction(s), credits {feats.get('total_in') or 0:,.0f} / "
            f"debits {feats.get('total_out') or 0:,.0f}" if txns else
            "no transactions attributed to this subject as account holder"
        )
        action = ("put forward for STR filing with FIU-IND, subject to analyst confirmation"
                  if r["band"] == "high" else
                  "flagged for analyst review; below the band this tool puts forward for "
                  "filing")
        lines.append(
            f"STR-{n:03d} | Suspected subject: {_str_subject(data, r)} | "
            f"Risk {r['risk_score']} ({r['band']}) | Grounds of suspicion: {grounds} | "
            f"Basis: {basis} | Particulars: {particulars} | Recommended action: {action}."
        )
    if not lines:
        return ["No entity reached the high or medium risk band with a fired typology, so no "
                "STR is drafted. This states that nothing matched — not that nothing was "
                "assessed; see the detection section for rules evaluated."]
    return lines
