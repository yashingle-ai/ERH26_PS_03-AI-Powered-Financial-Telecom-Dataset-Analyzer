"""Which narrative document mentions which identifier (FR-15 search over case paperwork).

150 Word documents across the two cases hold the investigation's own account of itself: bail
affidavits, charge sheets, arrest records, panchnamas, notices. Between them **36 carry their
evidence in prose only** and **8 are two-column key-value forms** — six of those the same 53-row
statutory bail affidavit, one per accused, whose single table holds 16,526 characters against 939
in its paragraphs. `docx_tables.py` reads tables as records, so a form whose header row is a
question and its answer maps to nothing, and prose is invisible to it entirely.

None of this can become events: an affidavit has no per-identifier timestamp. What it can do is
answer *"which documents mention this account?"* — the question an analyst actually asks when a
number surfaces and they need the narrative behind it. That is FR-15 applied to paperwork.

**These mentions are not merge keys and must never become them.** The documents are the officer's
case theory: an allegation in a legal filing, not a bank's record. `bank_reply_links` exists for
identity, precisely because a bank KYC reply is a different class of evidence — see
`docs/COMPONENT_STATUS.md` §4.2. A mention here says only *this document names this number*, which
is a pointer into the evidence and not a claim about who owns what.

The `(SECOND LAYER)` / `(THIRD LAYER)` annotations get the same treatment. They are the
investigating officer's own layering determination written beside the account it applies to —
valuable context, and an assertion, so they are carried as `asserted_layers` attributed to the
document rather than fed to the `layering` typology, which derives its own from the transfer graph.
"""

from __future__ import annotations

import re
from collections import Counter
from pathlib import Path

from ..core.logging_config import get_logger
from ..core.text import ascii_digits
from ..normalization import normalizers as nz

log = get_logger(__name__)

#: Identifier shapes. Deliberately anchored on a label where one exists, because an unanchored
#: 9-18 digit run in Gujarati prose is as likely to be a case number, a section citation or a
#: cheque serial as an account — the same reasoning that put AADHAAR behind keyword anchoring.
_LABELLED_ACCOUNT = re.compile(
    r"(?:A/?C(?:\s*No)?|ACCOUNT\s*(?:NO|NUMBER)?|એકાઉન્ટ\s*નં|ખાતા\s*નં|બેંક\s*એકાઉન્ટ)"
    r"[\s.:#–-]*([0-9]{9,18})", re.I)
_IFSC = re.compile(r"\b([A-Z]{4}0[A-Z0-9]{6})\b")
_MSISDN = re.compile(r"(?<!\d)([6-9]\d{9})(?!\d)")
_IMEI = re.compile(r"IMEI[^0-9]{0,12}((?:[0-9][0-9/\s,]{13,60}))", re.I)
_IMSI = re.compile(r"IMSI[^0-9]{0,12}([0-9]{15})", re.I)
_UPI = re.compile(r"\b([\w.\-]{3,}@(?:ok[a-z]+|[a-z]{2,}bank|paytm|ybl|axl|ibl|upi|apl|sbi))\b",
                  re.I)
#: The officer's own layering determination, written beside the account it applies to.
_LAYER = re.compile(r"\(\s*(FIRST|SECOND|THIRD|FOURTH|FIFTH|1ST|2ND|3RD|4TH)\s*LAYER\s*\)", re.I)

#: Document class, decided from its own text. Filenames are useless here — staging transliterates
#: Gujarati names into `એફ_ડ_વ_ટ`, so `affidavit` is unrecoverable from the path.
_DOC_KINDS = (
    ("affidavit", re.compile("એફીડેવીટ|એફિડેવિટ|affidavit", re.I)),
    ("chargesheet", re.compile("ચાર્જશીટ|charge.?sheet", re.I)),
    ("bail_application", re.compile("જામીન અરજી")),
    ("arrest_record", re.compile(r"\bMCR\b|ધરપકડ|અટક")),
    ("panchnama", re.compile("પંચનામ")),
    ("bank_reply", re.compile("બેંક એકાઉન્ટ નંબર")),
    ("notice", re.compile("યાદી|નોટીસ")),
    ("statement", re.compile("જવાબ|નિવેદન")),
)

#: A key-value form is a 2-column table whose left column is a short label. The bail affidavit is
#: 53 such rows; treating it as a record table produces one row of nonsense.
_KV_MAX_LABEL = 200
_KV_MIN_ROWS = 8


def _imei_digits(raw: str) -> str | None:
    """IMEIs are written `359788/09/777823/4` in these documents."""
    d = re.sub(r"\D", "", raw)
    return d if len(d) in (14, 15, 16) else None


def _classify(text: str) -> list[str]:
    kinds = [name for name, rx in _DOC_KINDS if rx.search(text)]
    return kinds or ["unclassified"]


def _extract(text: str) -> dict[str, list[str]]:
    text = ascii_digits(text)
    found: dict[str, set] = {
        "ACCOUNT_NO": {m for m in _LABELLED_ACCOUNT.findall(text)},
        "IFSC": set(_IFSC.findall(text)),
        "PHONE": {nz.phone(m) for m in _MSISDN.findall(text)},
        "IMEI": {d for d in (_imei_digits(m) for m in _IMEI.findall(text)) if d},
        "IMSI": set(_IMSI.findall(text)),
        "UPI_ID": {m.lower() for m in _UPI.findall(text)},
    }
    return {k: sorted(x for x in v if x) for k, v in found.items() if any(v)}


def _docx_parts(path: Path) -> tuple[str, str, int]:
    """(prose, table text, key-value table count). Kept apart because the split is the finding."""
    from docx import Document
    doc = Document(str(path))
    prose = "\n".join(p.text for p in doc.paragraphs)
    chunks, kv = [], 0
    for t in doc.tables:
        rows = t.rows
        if not rows:
            continue
        if len(t.columns) == 2 and len(rows) >= _KV_MIN_ROWS:
            labels = [r.cells[0].text.strip() for r in rows[:20]]
            if sum(1 for x in labels if 0 < len(x) <= _KV_MAX_LABEL) >= len(labels) * 0.6:
                kv += 1
        chunks.append("\n".join(c.text for r in rows for c in r.cells))
    return prose, "\n".join(chunks), kv


#: Every unreadable `.docx` across both cases is exactly this size — 47 of them.
_STUB_SIZE = 162


def _why_unreadable(path: Path, err: Exception) -> tuple[str, bool]:
    """A reason an analyst can act on, and whether the content was never delivered at all.

    All 47 unreadable `.docx` files across the two cases are **exactly 162 bytes**, each holding a
    length-prefixed source host name — `HP`, `ACER`, `pc`, `Admin`, `CYBER`, `admin`, six different
    machines. They are shortcut / cloud-placeholder stubs: the document's bytes were never copied
    into the evidence set. No reader can recover a file that was not delivered.

    The distinction matters because the two reasons ask for different things. "Unreadable" invites
    someone to try a better parser; "the content was never delivered" is a request to the case
    officer for the missing exhibit, and it belongs in the reject report in those words.
    """
    try:
        size = path.stat().st_size
        head = path.open("rb").read(16)
    except OSError:
        return f"could not be opened: {err}", False
    if size == _STUB_SIZE and not head.startswith(b"PK"):
        n = head[0] if head else 0
        host = head[1:1 + n].decode("latin-1", "replace") if 0 < n < 15 else "?"
        return (f"placeholder stub, not a document: {size} bytes naming source host "
                f"{host!r}. The file's content was never copied into the evidence set — "
                f"request the exhibit from the case officer"), True
    return f"unreadable as a Word document ({size} bytes): {err}", False


def build(input_dir: str, skipped_out: list | None = None) -> list[dict]:
    """One record per narrative document that mentions at least one identifier.

    `skipped_out` collects documents this indexer could not open, so they reach the reject report
    rather than living only in a log line. Five of the 155 `.docx` files across the two cases are
    unreadable by python-docx — the main ingestion path rejects them too, but a warning is not a
    reject, and rule 2 is that nothing is dropped silently.
    """
    root = Path(input_dir)
    if not root.exists():
        return []
    out: list[dict] = []
    for path in sorted(root.rglob("*.docx")):
        if not path.is_file() or path.name.startswith("~$"):
            continue
        try:
            prose, tables, kv_tables = _docx_parts(path)
        except Exception as e:
            reason, stub = _why_unreadable(path, e)
            log.warning("narrative document skipped %s: %s", path.name, reason)
            if skipped_out is not None:
                skipped_out.append({
                    "file": path.name, "path": str(path), "rows": 0, "rejected": 0,
                    "reason": reason, "stage": "document_mentions",
                    "evidentiary": True, "content_never_delivered": stub,
                })
            continue
        whole = f"{prose}\n{tables}"
        ids = _extract(whole)
        if not ids:
            continue
        out.append({
            "document": path.name,
            "relative_path": str(path.relative_to(root)),
            "kinds": _classify(whole),
            "identifiers": ids,
            "identifier_count": sum(len(v) for v in ids.values()),
            #: Where the evidence lives. `prose_only` documents are the ones a table reader
            #: cannot see at all, whatever the headers say.
            "prose_chars": len(prose),
            "table_chars": len(tables),
            "prose_only": bool(_extract(prose)) and not _extract(tables),
            "key_value_tables": kv_tables,
            #: The officer's assertion, attributed. NOT an input to the `layering` typology,
            #: which derives its own hops from the transfer graph.
            "asserted_layers": sorted({m.upper() for m in _LAYER.findall(whole)}),
        })
    if out:
        kinds = Counter(k for r in out for k in r["kinds"])
        log.info("indexed %d narrative documents mentioning identifiers from %s (%s)",
                 len(out), root.name,
                 ", ".join(f"{k}={v}" for k, v in kinds.most_common(5)))
    return out


def find(mentions: list[dict], value: str) -> list[dict]:
    """Documents mentioning `value`, matched as an identifier rather than as a substring.

    Substring search over a 16,000-character affidavit returns the document for any 4-digit run
    it happens to contain. The value is normalised the same way the index was, so a phone typed
    with spaces, a `+91`, or in Gujarati numerals finds the same document.
    """
    needle = ascii_digits(value).strip()
    if not needle:
        return []
    candidates = {needle, needle.upper(), needle.lower()}
    ph = nz.phone(needle)
    if ph:
        candidates.add(ph)
    acct = nz.account_no(needle)
    if acct:
        candidates.add(acct)
    hits = []
    for rec in mentions:
        matched = {kind: [v for v in vals if v in candidates]
                   for kind, vals in rec["identifiers"].items()}
        matched = {k: v for k, v in matched.items() if v}
        if matched:
            hits.append({**{k: rec[k] for k in
                            ("document", "relative_path", "kinds", "asserted_layers")},
                         "matched": matched})
    return hits
