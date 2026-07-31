"""Narrative case paperwork, indexed by the identifiers it names (FR-15).

36 of the 150 Word documents in the two cases carry their evidence in prose only, and 8 are
two-column key-value forms — six of those the same 53-row bail affidavit, whose single table holds
16,526 characters against 939 in its paragraphs. A record-table reader maps a form whose header row
is a question to nothing, and cannot see prose at all.

The two things this must get right are the reason it exists:

  * **Match identifiers, not substrings.** A substring search over a 16,000-character affidavit
    returns the document for any four-digit run it contains.
  * **Never become a merge key.** These documents are the officer's case theory. Identity comes from
    `bank_reply_links`, because a bank's KYC reply is a different class of evidence.

All fixtures are **synthetic**.
"""

from __future__ import annotations

from docx import Document

from backend.app.search import document_mentions as dm

G = str.maketrans("0123456789", "૦૧૨૩૪૫૬૭૮૯")


def _doc(tmp_path, name, paragraphs=(), kv_rows=(), table_rows=()):
    d = Document()
    for p in paragraphs:
        d.add_paragraph(p)
    if kv_rows:
        t = d.add_table(rows=0, cols=2)
        for label, value in kv_rows:
            c = t.add_row().cells
            c[0].text, c[1].text = label, value
    if table_rows:
        t2 = d.add_table(rows=0, cols=len(table_rows[0]))
        for row in table_rows:
            c = t2.add_row().cells
            for i, v in enumerate(row):
                c[i].text = str(v)
    path = tmp_path / name
    d.save(str(path))
    return path


# ── extraction ──────────────────────────────────────────────────────────────────────

def test_a_labelled_account_in_prose_is_indexed(tmp_path):
    _doc(tmp_path, "affidavit.docx",
         paragraphs=["આ કામે એફીડેવીટ કરેલ છે.",
                     "UCO BANK ના એકાઉન્ટ નં. 234802100017 (IFSC:UCDA0002348) નો ધારક"])
    recs = dm.build(str(tmp_path))
    assert len(recs) == 1
    ids = recs[0]["identifiers"]
    assert "234802100017" in ids["ACCOUNT_NO"]
    assert "UCDA0002348" in ids["IFSC"]
    assert "affidavit" in recs[0]["kinds"]


def test_an_unlabelled_digit_run_is_not_claimed_as_an_account(tmp_path):
    """An unanchored 9-18 digit run in Gujarati prose is as likely to be a case number or a
    section citation as an account — the same reasoning that put AADHAAR behind anchoring."""
    _doc(tmp_path, "notice.docx",
         paragraphs=["યાદી", "ગુ.ર.નં. 11210062240065 તથા સી.સી.નં. 154672025 મુજબ",
                     "એકાઉન્ટ નં. 987654321012 નો ધારક"])
    ids = dm.build(str(tmp_path))[0]["identifiers"]
    assert ids["ACCOUNT_NO"] == ["987654321012"], "only the labelled one"


def test_a_slashed_imei_and_an_imsi_are_recovered(tmp_path):
    """These documents write `IMEI No.359788/09/777823/4` and `IMSI નં. 404222454920891`."""
    _doc(tmp_path, "panchnama.docx",
         paragraphs=["પંચનામુ", "IMEI No.359788/09/777823/4 તથા IMSI નં. 404222454920891"])
    ids = dm.build(str(tmp_path))[0]["identifiers"]
    assert "359788097778234" in ids["IMEI"]
    assert "404222454920891" in ids["IMSI"]


def test_gujarati_numerals_are_indexed_as_ascii(tmp_path):
    """~60 MSISDNs across the corpus are written in Gujarati numerals. Nothing could match them
    before `core.text.ascii_digits`."""
    _doc(tmp_path, "yadi.docx", paragraphs=["યાદી", "મો.નં." + "9824444401".translate(G)])
    ids = dm.build(str(tmp_path))[0]["identifiers"]
    assert ids["PHONE"] == ["+919824444401"]


def test_key_value_tables_are_counted_and_their_content_read(tmp_path):
    """The bail affidavit is 53 rows of `(૧) label | value`. Read as a record table it maps to
    nothing, so the content has to come from the cells regardless of the shape."""
    rows = [(f"({i}) વિગત", f"detail {i}") for i in range(1, 12)]
    rows.append(("(૧૨) બેંક", "HDFC A/C No. 502000994124 (IFSC:HDFC0001234)"))
    _doc(tmp_path, "efidevit.docx", paragraphs=["એફીડેવીટ"], kv_rows=rows)
    rec = dm.build(str(tmp_path))[0]
    assert rec["key_value_tables"] == 1
    assert "502000994124" in rec["identifiers"]["ACCOUNT_NO"]


def test_a_prose_only_document_is_flagged_as_such(tmp_path):
    """These are the ones a table reader cannot see whatever the headers say."""
    _doc(tmp_path, "prose.docx", paragraphs=["જવાબ", "એકાઉન્ટ નં. 111222333444 નો ધારક"])
    assert dm.build(str(tmp_path))[0]["prose_only"] is True


def test_a_document_with_no_identifiers_is_not_indexed(tmp_path):
    _doc(tmp_path, "empty.docx", paragraphs=["રૂબરૂ", "પોલીસ ઇન્સ્પેક્ટર"])
    assert dm.build(str(tmp_path)) == []


def test_layer_annotations_are_carried_as_assertions(tmp_path):
    """The officer's own layering determination. Attributed context, never an input to the
    `layering` typology, which derives its hops from the transfer graph."""
    _doc(tmp_path, "cs.docx",
         paragraphs=["ચાર્જશીટ",
                     "CITY UNION BANK A/C : 500101014008 (THIRD LAYER)",
                     "YES BANK A/C : 040061900017 (SECOND LAYER) ના ધારક"])
    rec = dm.build(str(tmp_path))[0]
    assert rec["asserted_layers"] == ["SECOND", "THIRD"]


def test_documents_are_classified_from_their_text_not_their_filename(tmp_path):
    """Staging transliterates Gujarati filenames into `એફ_ડ_વ_ટ`, so the path cannot be trusted."""
    _doc(tmp_path, "0065-2024__AROPI__x_y_z.docx",
         paragraphs=["એફીડેવીટ", "જામીન અરજી નં. 2176/2025", "એકાઉન્ટ નં. 123456789012"])
    kinds = dm.build(str(tmp_path))[0]["kinds"]
    assert "affidavit" in kinds and "bail_application" in kinds


def test_an_unreadable_document_does_not_stop_the_others(tmp_path):
    (tmp_path / "broken.docx").write_bytes(b"not a docx")
    _doc(tmp_path, "ok.docx", paragraphs=["એકાઉન્ટ નં. 555666777888"])
    assert len(dm.build(str(tmp_path))) == 1


def test_a_missing_folder_is_not_an_error(tmp_path):
    assert dm.build(str(tmp_path / "nope")) == []


# ── lookup ──────────────────────────────────────────────────────────────────────────

def _indexed(tmp_path):
    _doc(tmp_path, "a.docx", paragraphs=["એફીડેવીટ", "એકાઉન્ટ નં. 234802100017", "મો.નં.9824444401"])
    _doc(tmp_path, "b.docx", paragraphs=["ચાર્જશીટ", "એકાઉન્ટ નં. 999888777666"])
    return dm.build(str(tmp_path))


def test_find_returns_only_documents_naming_that_identifier(tmp_path):
    recs = _indexed(tmp_path)
    hits = dm.find(recs, "234802100017")
    assert [h["document"] for h in hits] == ["a.docx"]
    assert hits[0]["matched"]["ACCOUNT_NO"] == ["234802100017"]


def test_find_matches_an_identifier_not_a_substring(tmp_path):
    """`2348` occurs inside the account number. A substring search would return the document."""
    assert dm.find(_indexed(tmp_path), "2348") == []


def test_find_normalises_the_needle_the_same_way_as_the_index(tmp_path):
    recs = _indexed(tmp_path)
    for form in ("9824444401", "+919824444401", "98244 44401",
                 "9824444401".translate(G)):
        assert [h["document"] for h in dm.find(recs, form)] == ["a.docx"], form


def test_find_on_an_absent_or_empty_value_returns_nothing(tmp_path):
    recs = _indexed(tmp_path)
    assert dm.find(recs, "000000000000") == []
    assert dm.find(recs, "") == []
    assert dm.find(recs, "   ") == []


def test_a_mention_carries_no_identity_claim(tmp_path):
    """The guard that matters. A record must expose the document and what it names, and nothing
    that entity resolution could consume as a merge — identity comes from `bank_reply_links`."""
    rec = _indexed(tmp_path)[0]
    assert "own_identifiers" not in rec, "that key is what entity resolution merges on"
    assert "primary" not in rec
    assert rec.get("event_type") is None


def test_an_unreadable_document_reaches_the_reject_report(tmp_path):
    """A warning in a log is not a reject. Five of the 155 .docx across the two cases cannot be
    opened by python-docx, and rule 2 is that nothing is dropped silently."""
    (tmp_path / "broken.docx").write_bytes(b"not a docx")
    _doc(tmp_path, "ok.docx", paragraphs=["એકાઉન્ટ નં. 555666777888"])
    skips: list = []
    recs = dm.build(str(tmp_path), skipped_out=skips)
    assert len(recs) == 1
    assert len(skips) == 1
    assert skips[0]["file"] == "broken.docx"
    assert "unreadable" in skips[0]["reason"]
    assert skips[0]["stage"] == "document_mentions"


def test_the_pipeline_puts_those_skips_in_the_reject_report(tmp_path):
    """The indexer can collect them and the pipeline can still drop them on the floor — which is
    the shape F1 and F3 were both in."""
    import inspect

    from backend.app import pipeline
    src = inspect.getsource(pipeline.run_base)
    assert "skipped_out=mention_skips" in src
    assert "inv.rejects += mention_skips" in src


def test_a_placeholder_stub_is_named_as_undelivered_not_as_unreadable(tmp_path):
    """All 47 unreadable .docx across the two cases are exactly 162 bytes holding a
    length-prefixed source host name — shortcut/cloud-placeholder stubs whose content was never
    copied into the evidence set. "Unreadable" invites a better parser; "never delivered" is a
    request to the case officer, and that difference belongs in the reject report."""
    stub = bytes([2]) + b"HP\x00" + b"\x00" * (162 - 4)
    (tmp_path / "missing.docx").write_bytes(stub)
    skips: list = []
    dm.build(str(tmp_path), skipped_out=skips)
    assert len(skips) == 1
    assert skips[0]["content_never_delivered"] is True
    assert "never copied into the evidence set" in skips[0]["reason"]
    assert "'HP'" in skips[0]["reason"]


def test_a_genuinely_corrupt_document_is_not_called_undelivered(tmp_path):
    """A wrong claim in the other direction would send the officer looking for a file we have."""
    (tmp_path / "corrupt.docx").write_bytes(b"PK\x03\x04" + b"garbage" * 40)
    skips: list = []
    dm.build(str(tmp_path), skipped_out=skips)
    assert len(skips) == 1
    assert skips[0]["content_never_delivered"] is False
    assert "unreadable as a Word document" in skips[0]["reason"]
